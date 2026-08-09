from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from .metric_names import canonical_metric, scoped_metric_name


TEXT_SUFFIXES = {".md", ".txt", ".rst"}
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_PDF_PAGES = 1_000
_NUMERIC_STRING_RE = re.compile(r"[+-]?(?:\d+(?:[.,]\d*)?|[.,]\d+)(?:[eE][+-]?\d+)?")


def extract_document_text(path: Path, selector: str | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        try:
            with zipfile.ZipFile(path) as archive:
                expanded_size = sum(item.file_size for item in archive.infolist())
        except zipfile.BadZipFile as error:
            raise ValueError("report DOCX is not a valid ZIP container") from error
        if expanded_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
            raise ValueError("report DOCX expands beyond the 100 MB safety limit")
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(paragraphs)
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(path)
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError("report PDF exceeds the 1000-page safety limit")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".ipynb":
        payload = json.loads(path.read_text(encoding="utf-8"))
        return "\n".join(
            _source_text(cell.get("source", ""))
            for cell in payload.get("cells", [])
            if cell.get("cell_type") == "markdown"
        )
    if suffix == ".json":
        payload = json.loads(path.read_text(encoding="utf-8"))
        selected = _select(payload, selector) if selector else payload
        return _json_text(selected)
    raise ValueError("report must be Markdown, TXT, RST, DOCX, PDF, IPYNB, or JSON")


def _source_text(source: str | list[str]) -> str:
    return "".join(source) if isinstance(source, list) else source


def _select(payload: object, selector: str) -> object:
    current = payload
    for part in selector.split("."):
        if isinstance(current, list):
            try:
                current = current[int(part)]
            except (ValueError, IndexError) as error:
                raise ValueError(f"invalid JSON selector: {selector}") from error
        elif isinstance(current, dict) and part in current:
            current = current[part]
        else:
            raise ValueError(f"invalid JSON selector: {selector}")
    return current


def _json_text(value: object, path: tuple[object, ...] = ()) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        lines = []
        for key, item in value.items():
            item_path = (*path, key)
            if isinstance(item, (int, float)) and not isinstance(item, bool):
                name = scoped_metric_name(item_path) or canonical_metric(key)
                lines.append(f"{name}: {item}")
            elif (
                isinstance(item, str)
                and scoped_metric_name(item_path) is not None
                and _NUMERIC_STRING_RE.fullmatch(item.strip())
            ):
                lines.append(f"{scoped_metric_name(item_path)}: {item.strip()}")
            else:
                lines.append(_json_text(item, item_path))
        return "\n".join(line for line in lines if line)
    if isinstance(value, list):
        return "\n".join(_json_text(item, path) for item in value)
    return str(value)
