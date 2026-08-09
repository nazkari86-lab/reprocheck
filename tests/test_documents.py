import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from reprocheck import documents
from reprocheck.documents import extract_document_text


def test_extracts_docx_paragraphs_and_tables(tmp_path: Path):
    path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Accuracy: 91%")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "F1"
    table.cell(0, 1).text = "0.90"
    document.save(str(path))
    text = extract_document_text(path)
    assert "Accuracy: 91%" in text
    assert "F1 | 0.90" in text


def test_extracts_markdown_cells_from_notebook(tmp_path: Path):
    path = tmp_path / "report.ipynb"
    path.write_text(
        json.dumps(
            {
                "cells": [
                    {"cell_type": "markdown", "source": ["Accuracy: ", "95%"]},
                    {"cell_type": "code", "source": "print('ignored')"},
                ]
            }
        ),
        encoding="utf-8",
    )
    assert extract_document_text(path) == "Accuracy: 95%"


def test_extracts_textual_notebook_outputs_without_executing_code(tmp_path: Path):
    path = tmp_path / "outputs.ipynb"
    path.write_text(
        json.dumps(
            {
                "cells": [
                    {
                        "cell_type": "code",
                        "source": "raise RuntimeError('must not execute')",
                        "outputs": [
                            {"output_type": "stream", "text": ["Accuracy: 91%\n"]},
                            {
                                "output_type": "execute_result",
                                "data": {"text/plain": ["F1: 0.88"]},
                            },
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    assert extract_document_text(path) == "Accuracy: 91%\n\nF1: 0.88"


def test_rejects_malformed_notebook_report_contract(tmp_path: Path):
    path = tmp_path / "bad.ipynb"
    path.write_text('{"cells": "not-an-array"}', encoding="utf-8")
    with pytest.raises(ValueError, match="cells array"):
        extract_document_text(path)


def test_extracts_selected_json_claim(tmp_path: Path):
    path = tmp_path / "claims.json"
    path.write_text(
        json.dumps({"claims": [{"claim": "Hard Dice 0.9036 and hard IoU 0.8242"}]}),
        encoding="utf-8",
    )
    assert extract_document_text(path, "claims.0.claim") == ("Hard Dice 0.9036 and hard IoU 0.8242")


def test_preserves_numeric_json_keys_as_claim_context(tmp_path: Path):
    path = tmp_path / "metadata.json"
    path.write_text(
        json.dumps(
            {
                "name": "public model",
                "eval_metrics": {
                    "val_mean_dice": 0.8518,
                    "validation_accuracy": 94.5,
                    "mean_dice": {"central gland": 0.88, "peripheral zone": 0.75},
                },
            }
        ),
        encoding="utf-8",
    )
    text = extract_document_text(path)
    assert "val_mean_dice: 0.8518" in text
    assert "validation_accuracy: 94.5" in text
    assert "mean_dice_central_gland: 0.88" in text
    assert "mean_dice_peripheral_zone: 0.75" in text


def test_preserves_numeric_strings_only_when_the_json_path_is_a_metric(tmp_path: Path):
    path = tmp_path / "metadata.json"
    path.write_text(
        json.dumps(
            {
                "eval_metrics": {"validation_accuracy": "0.95"},
                "run_id": "12345",
            }
        ),
        encoding="utf-8",
    )
    text = extract_document_text(path)
    assert "validation_accuracy: 0.95" in text
    assert "12345" in text
    assert "run_id: 12345" not in text


def test_rejects_docx_expansion_over_safety_limit(tmp_path: Path, monkeypatch):
    path = tmp_path / "report.docx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "x" * 11)
    monkeypatch.setattr(documents, "MAX_ARCHIVE_UNCOMPRESSED_BYTES", 10)
    with pytest.raises(ValueError, match="safety limit"):
        extract_document_text(path)


def test_rejects_invalid_docx_container(tmp_path: Path):
    path = tmp_path / "report.docx"
    path.write_bytes(b"not-a-zip")
    with pytest.raises(ValueError, match="valid ZIP"):
        extract_document_text(path)


def test_pdf_page_limit_is_enforced(tmp_path: Path, monkeypatch):
    path = tmp_path / "report.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as handle:
        writer.write(handle)
    monkeypatch.setattr(documents, "MAX_PDF_PAGES", 0)
    with pytest.raises(ValueError, match="page safety limit"):
        extract_document_text(path)


def test_json_selector_and_suffix_errors_are_explicit(tmp_path: Path):
    path = tmp_path / "report.json"
    path.write_text('{"claims":["Accuracy: 90%"]}', encoding="utf-8")
    assert extract_document_text(path) == "Accuracy: 90%"
    with pytest.raises(ValueError, match="invalid JSON selector"):
        extract_document_text(path, "claims.bad")

    unsupported = tmp_path / "report.yaml"
    unsupported.write_text("accuracy: 0.9", encoding="utf-8")
    with pytest.raises(ValueError, match="report must be"):
        extract_document_text(unsupported)
